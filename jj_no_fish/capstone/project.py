from io import BytesIO
from docx import Document
from django.core.files.base import ContentFile
from django.db import models
import json
from .models import UploadedDocument

def main(action, request, word_pairs=None, user=None):
    if action == 'Upload & Replace' and 'document' in request.FILES:
        # Fix the parameter order here to match the function definition
        uploaded_doc = handle_uploaded_document(request.FILES['document'], word_pairs, user)
        return uploaded_doc

    elif action == 'Export':
        uploaded_doc = export_modified_document()
        return uploaded_doc

    elif action == 'jj_no_fish':
        # Return default context without upload_completed flag
        return {
            'old_words': [],
            'new_words': []
        }

    return {}

def replace_words_in_document(document_file, word_pairs):
    doc = Document(document_file)
    original_text = []
    modified_text = []

    # First collect original text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
            original_text.append(paragraph.text)

    # Then process and collect modified text
    doc = Document(document_file)  # Create fresh document for modifications
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if word_pairs:
            for old_word, new_word in word_pairs.items():
                if old_word and new_word:
                    text = text.replace(old_word, new_word)
        paragraph.text = text
        if text.strip():  # Only add non-empty paragraphs
            modified_text.append(text)

    modified_document = BytesIO()
    doc.save(modified_document)
    modified_document.seek(0)

    return {
        'document': modified_document.getvalue(),
        'original_preview': '\n'.join(original_text[:10]),  # First 10 paragraphs
        'modified_preview': '\n'.join(modified_text[:10])
    }

def handle_uploaded_document(file, word_pairs, user):
    new_document = UploadedDocument(document=file)
    new_document.user = user
    new_document.save()

    result = replace_words_in_document(new_document.document, word_pairs)
    new_document.modified_document.save(f'{file.name}_modified.docx', ContentFile(result['document']))
    new_document.word_pairs = json.dumps(word_pairs)
    new_document.save()

    return {
        'upload_completed': True,
        'word_pairs': word_pairs,
        'modified_doc': result['document'],
        'original_preview': result['original_preview'],
        'modified_preview': result['modified_preview'] if word_pairs else result['original_preview'],
        'new_document_id': new_document.id
    }

def export_modified_document():
    uploaded_doc = UploadedDocument.objects.last()
    if uploaded_doc and uploaded_doc.modified_document:
        file_name = uploaded_doc.modified_document.name
        file_name = file_name.replace("modified_document_","")
        return {'modified_doc': uploaded_doc.modified_document.read(), 'new_file_name': file_name}
    else:
        return {}